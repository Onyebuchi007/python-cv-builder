from docx import Document
from docx.shared import Inches
import pyttsx3


def talk(text):
    pyttsx3.speak(text)


document = Document()

# Profile pic
document.add_picture('Buchi.jpeg', width=Inches(2.0))

# Bio
name = input('What is your name? ')
talk(f'Hello {name}, how are you today?')
talk('What is your phone number?')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

document.add_paragraph(
    name + '|' + phone_number + '|' + email
)

# Profile
document.add_heading('About me')

about_me = input('Tell us about yourself? ')
document.add_paragraph(about_me)

# Work Experience
document.add_heading('Work Experiences')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')
experience_details = input(f'Describe your experience at {company} ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-').italic = True
p.add_run(to_date + '\n').italic = True
p.add_run(experience_details)

# More Experiences
while True:
    more_experiences = input('Do you have more work experiences? y for yes, n for no ')
    if more_experiences.lower() == 'y':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')
        experience_details = input(f'Describe your experience at {company} ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-').italic = True
        p.add_run(to_date + '\n').italic = True
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
skill = input('Add a skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# More Skills
while True:
    more_skills = input('Do you have additional skills to add? y for yes and n for no ')
    if more_skills.lower() == 'y':
        skill = input('Add a skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated via Tutorial, credit - Amigoscode'

document.save('cv2.docx')
